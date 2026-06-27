from __future__ import annotations

import io
from datetime import datetime, timezone

from docx import Document

from .models import Project, ReviewTask


def build_response_letter_docx(project: Project, tasks: list[ReviewTask]) -> bytes:
    approved_tasks = [task for task in tasks if task.status == "Approved"]
    if not approved_tasks:
        raise ValueError("At least one approved task is required before export.")

    doc = Document()
    doc.add_heading("Response to Reviewers", 0)
    doc.add_paragraph(f"Manuscript: {project.title}")
    doc.add_paragraph(f"Journal: {project.journal}")
    doc.add_paragraph(f"Prepared: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}")
    doc.add_paragraph(
        "We sincerely thank the reviewers for their careful evaluation and constructive comments. "
        "Our approved point-by-point responses are provided below."
    )

    for task in approved_tasks:
        doc.add_heading(f"{task.id} - {task.title}", level=2)
        doc.add_paragraph("Comment", style="Heading 3")
        doc.add_paragraph(task.comment)
        doc.add_paragraph("Response", style="Heading 3")
        doc.add_paragraph(task.response_draft)
        doc.add_paragraph("Change in manuscript", style="Heading 3")
        doc.add_paragraph(f"{task.manuscript_section}: {task.suggested_change}")

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()
